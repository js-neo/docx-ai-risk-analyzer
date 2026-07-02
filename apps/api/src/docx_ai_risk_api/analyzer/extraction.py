from io import BytesIO

from docx import Document


class DocumentExtractionError(Exception):
    pass


def extract_text_from_docx_bytes(content: bytes) -> str:
    try:
        document = Document(BytesIO(content))
    except Exception as exc:
        raise DocumentExtractionError("Failed to read DOCX document") from exc

    blocks: list[str] = []

    for paragraph in document.paragraphs:
        text = paragraph.text.strip()
        if text:
            blocks.append(text)

    for table in document.tables:
        for row in table.rows:
            cells: list[str] = []

            for cell in row.cells:
                cell_text = " ".join(
                    paragraph.text.strip()
                    for paragraph in cell.paragraphs
                    if paragraph.text.strip()
                )

                if cell_text:
                    cells.append(cell_text)

            if cells:
                blocks.append(" | ".join(cells))

    text = "\n\n".join(blocks).strip()

    if not text:
        raise DocumentExtractionError("DOCX document does not contain readable text")

    return text
